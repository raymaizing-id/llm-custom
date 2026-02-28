#!/usr/bin/env python3
"""Generate AI service business analysis (Markdown, CSV, DOCX, charts) from web-scraped public data."""

from __future__ import annotations

import json
import math
import re
import time
from dataclasses import dataclass
from datetime import UTC, datetime, timedelta
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
import pandas as pd
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches


OUT_DIR = Path(
    "/home/raymaizing/Bisnis/workspace/00-second-brain/brain-raymaizing/00-inbox"
)
CSV_PATH = OUT_DIR / "ai_jasa_data_analytics.csv"
MD_PATH = OUT_DIR / "bisnis-jasa-ai-paling-potensial.md"
DOCX_PATH = OUT_DIR / "ai_jasa_data_analytics.docx"
SCRAPE_LOG_PATH = OUT_DIR / "ai_jasa_scrape_log.csv"
CHART_SCORE = OUT_DIR / "chart_opportunity_score.png"
CHART_SCATTER = OUT_DIR / "chart_cagr_vs_buzz.png"
CHART_HEATMAP = OUT_DIR / "chart_signal_heatmap.png"


@dataclass(frozen=True)
class SectorConfig:
    sector: str
    slug: str
    github_query: str
    hn_query: str
    arxiv_query: str
    team_fit_score: float
    service_pack: str
    target_client: str
    pricing_model: str


SECTORS: list[SectorConfig] = [
    SectorConfig(
        sector="MLaaS Platform & MLOps Enablement",
        slug="machine-learning-as-a-service-market",
        github_query="mlops",
        hn_query="machine learning as a service",
        arxiv_query='"machine learning as a service"',
        team_fit_score=8.9,
        service_pack="Platform MLOps, deployment model, observability, governance, dan retraining pipeline.",
        target_client="Startup SaaS, fintech, dan enterprise mid-market.",
        pricing_model="Setup fee + retainer bulanan + usage model endpoint.",
    ),
    SectorConfig(
        sector="Generative AI Integration Service",
        slug="generative-ai-market",
        github_query="generative ai",
        hn_query="generative ai integration",
        arxiv_query='"generative ai"',
        team_fit_score=9.1,
        service_pack="Integrasi LLM ke workflow bisnis, RAG, guardrails, dan evaluasi kualitas.",
        target_client="Perusahaan dengan knowledge base besar dan proses manual dokumen.",
        pricing_model="Implementation fee + subscription support + token usage pass-through.",
    ),
    SectorConfig(
        sector="Intelligent Document Processing (IDP) Service",
        slug="intelligent-document-processing-market",
        github_query="intelligent document processing",
        hn_query="intelligent document processing",
        arxiv_query='"intelligent document processing"',
        team_fit_score=8.7,
        service_pack="OCR + extraction + validation + workflow approval untuk invoice/kontrak/form.",
        target_client="BPO, legal ops, akuntansi, dan perusahaan logistik.",
        pricing_model="Biaya implementasi + harga per dokumen diproses.",
    ),
    SectorConfig(
        sector="Conversational AI Customer Automation",
        slug="conversational-ai-market",
        github_query="conversational ai chatbot",
        hn_query="conversational ai customer support",
        arxiv_query='"conversational ai"',
        team_fit_score=8.8,
        service_pack="Chatbot/voicebot omnichannel, integrasi CRM/helpdesk, dan analytics CSAT.",
        target_client="E-commerce, healthcare, dan layanan pelanggan volume tinggi.",
        pricing_model="Setup + retainer + tier berdasarkan volume percakapan.",
    ),
    SectorConfig(
        sector="Predictive Analytics Service",
        slug="predictive-analytics-market",
        github_query="predictive analytics",
        hn_query="predictive analytics service",
        arxiv_query='"predictive analytics"',
        team_fit_score=8.3,
        service_pack="Forecasting demand/churn/risk, dashboard eksekutif, dan alert otomatis.",
        target_client="Retail, D2C, dan operation-heavy businesses.",
        pricing_model="Project-based + lisensi dashboard + maintenance bulanan.",
    ),
    SectorConfig(
        sector="RPA + AI Workflow Automation Service",
        slug="robotic-process-automation-market",
        github_query="robotic process automation",
        hn_query="robotic process automation ai",
        arxiv_query='"robotic process automation"',
        team_fit_score=8.5,
        service_pack="Automasi proses back-office terintegrasi AI untuk triase, routing, dan QA.",
        target_client="Perusahaan dengan proses berulang lintas sistem legacy.",
        pricing_model="Audit proses + biaya implementasi per workflow + SLA support.",
    ),
]


def safe_float(value: str) -> float:
    return float(value.replace(",", "").strip())


def minmax(series: pd.Series) -> pd.Series:
    s = series.astype(float)
    lo, hi = s.min(), s.max()
    if math.isclose(lo, hi):
        return pd.Series([0.5] * len(s), index=s.index)
    return (s - lo) / (hi - lo)


def fetch_json(
    session: requests.Session, url: str, *, params: dict[str, Any] | None = None
) -> tuple[dict[str, Any], int, int]:
    last_error = "unknown"
    for attempt in range(1, 4):
        try:
            resp = session.get(url, params=params, timeout=30)
            status = resp.status_code
            if status == 200:
                return resp.json(), status, attempt

            # Respect GitHub rate-limit reset when needed.
            if status == 403 and "api.github.com" in url:
                reset = resp.headers.get("X-RateLimit-Reset")
                if reset and reset.isdigit():
                    wait_seconds = max(1, int(reset) - int(time.time()) + 1)
                    time.sleep(min(wait_seconds, 40))
                else:
                    time.sleep(2 * attempt)
            else:
                time.sleep(1.5 * attempt)
            last_error = f"http_{status}"
        except Exception as exc:  # noqa: BLE001
            last_error = str(exc)
            time.sleep(1.5 * attempt)
    raise RuntimeError(f"Failed JSON fetch after 3 attempts: {url} ({last_error})")


def fetch_text(session: requests.Session, url: str) -> tuple[str, int, int]:
    last_error = "unknown"
    for attempt in range(1, 4):
        try:
            resp = session.get(url, timeout=30)
            if resp.status_code == 200:
                return resp.text, resp.status_code, attempt
            last_error = f"http_{resp.status_code}"
            time.sleep(1.5 * attempt)
        except Exception as exc:  # noqa: BLE001
            last_error = str(exc)
            time.sleep(1.5 * attempt)
    raise RuntimeError(f"Failed text fetch after 3 attempts: {url} ({last_error})")


def parse_market_meta(html: str) -> dict[str, Any]:
    soup = BeautifulSoup(html, "lxml")
    title = soup.title.text.strip() if soup.title else ""
    meta_tag = soup.find("meta", attrs={"name": "description"})
    description = meta_tag["content"].strip() if meta_tag and meta_tag.get("content") else ""
    description = re.sub(r"\s+", " ", description)

    if "Precedence Research and Consulting:" in description:
        raise ValueError("Generic page description; likely redirected to homepage.")

    cagr_match = re.search(r"CAGR(?:\s+of)?\s+([0-9]+(?:\.[0-9]+)?)%", description, re.I)
    cagr = safe_float(cagr_match.group(1)) if cagr_match else None

    base_match = re.search(
        r"USD\s*([0-9,]+(?:\.[0-9]+)?)\s*(?:billion|bn)[^.]{0,100}?\bin\s*(20\d{2})",
        description,
        re.I,
    )
    target_match = re.search(
        r"USD\s*([0-9,]+(?:\.[0-9]+)?)\s*(?:billion|bn)[^.]{0,100}?\bby\s*(20\d{2})",
        description,
        re.I,
    )

    if not (base_match and target_match and cagr is not None):
        raise ValueError("Could not parse market base/target/CAGR from description.")

    base_value, base_year = safe_float(base_match.group(1)), int(base_match.group(2))
    target_value, target_year = safe_float(target_match.group(1)), int(target_match.group(2))
    implied_cagr = (pow(target_value / base_value, 1 / (target_year - base_year)) - 1) * 100

    return {
        "title": title,
        "description": description,
        "market_base_usd_billion": base_value,
        "market_base_year": base_year,
        "market_target_usd_billion": target_value,
        "market_target_year": target_year,
        "market_cagr_pct": cagr,
        "implied_cagr_pct": implied_cagr,
    }


def build_dataset() -> tuple[pd.DataFrame, pd.DataFrame]:
    session = requests.Session()
    session.headers.update(
        {
            "User-Agent": (
                "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
                "(KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36 CodexAgent/1.0"
            )
        }
    )
    scrape_log: list[dict[str, Any]] = []
    rows: list[dict[str, Any]] = []

    for sector in SECTORS:
        row: dict[str, Any] = {
            "sector": sector.sector,
            "market_source_url": f"https://www.precedenceresearch.com/{sector.slug}",
            "github_query": sector.github_query,
            "hn_query": sector.hn_query,
            "arxiv_query": sector.arxiv_query,
            "team_fit_score": sector.team_fit_score,
            "recommended_service_pack": sector.service_pack,
            "target_client": sector.target_client,
            "pricing_model": sector.pricing_model,
        }

        # 1) Market report scrape
        market_text, market_status, market_attempt = fetch_text(session, row["market_source_url"])
        market_data = parse_market_meta(market_text)
        row.update(market_data)
        scrape_log.append(
            {
                "sector": sector.sector,
                "source_type": "market_report",
                "url": row["market_source_url"],
                "status_code": market_status,
                "attempt_used": market_attempt,
                "captured_at": datetime.now().isoformat(timespec="seconds"),
            }
        )

        # 2) GitHub signal (developer supply / competition intensity)
        gh_json, gh_status, gh_attempt = fetch_json(
            session,
            "https://api.github.com/search/repositories",
            params={"q": sector.github_query, "per_page": 1},
        )
        row["github_repo_count"] = int(gh_json.get("total_count", 0))
        scrape_log.append(
            {
                "sector": sector.sector,
                "source_type": "github_search",
                "url": "https://api.github.com/search/repositories",
                "status_code": gh_status,
                "attempt_used": gh_attempt,
                "captured_at": datetime.now().isoformat(timespec="seconds"),
            }
        )

        # 3) Hacker News signal (market buzz from founders/builders)
        one_year_ago = int((datetime.now(UTC) - timedelta(days=365)).timestamp())
        hn_json, hn_status, hn_attempt = fetch_json(
            session,
            "https://hn.algolia.com/api/v1/search_by_date",
            params={
                "query": sector.hn_query,
                "tags": "story",
                "numericFilters": f"created_at_i>{one_year_ago}",
            },
        )
        row["hn_story_hits_last_12m"] = int(hn_json.get("nbHits", 0))
        scrape_log.append(
            {
                "sector": sector.sector,
                "source_type": "hn_algolia",
                "url": "https://hn.algolia.com/api/v1/search_by_date",
                "status_code": hn_status,
                "attempt_used": hn_attempt,
                "captured_at": datetime.now().isoformat(timespec="seconds"),
            }
        )

        # 4) arXiv signal (innovation momentum)
        arxiv_text, arxiv_status, arxiv_attempt = fetch_text(
            session,
            "https://export.arxiv.org/api/query"
            f"?search_query=all:{requests.utils.quote(sector.arxiv_query)}&start=0&max_results=1",
        )
        total_match = re.search(
            r"<opensearch:totalResults>(\d+)</opensearch:totalResults>", arxiv_text
        )
        row["arxiv_total_results"] = int(total_match.group(1)) if total_match else 0
        scrape_log.append(
            {
                "sector": sector.sector,
                "source_type": "arxiv_api",
                "url": "https://export.arxiv.org/api/query",
                "status_code": arxiv_status,
                "attempt_used": arxiv_attempt,
                "captured_at": datetime.now().isoformat(timespec="seconds"),
            }
        )

        rows.append(row)

    df = pd.DataFrame(rows)
    log_df = pd.DataFrame(scrape_log)

    # Derived metrics and opportunity score.
    df["growth_multiple_x"] = df["market_target_usd_billion"] / df["market_base_usd_billion"]
    df["cagr_norm"] = minmax(df["market_cagr_pct"])
    df["target_market_norm"] = minmax(df["market_target_usd_billion"])
    df["hn_norm"] = minmax(df["hn_story_hits_last_12m"])
    df["arxiv_norm"] = minmax(df["arxiv_total_results"])
    df["competition_norm"] = minmax(df["github_repo_count"])
    df["fit_norm"] = df["team_fit_score"] / 10.0

    df["demand_signal_score"] = (
        0.40 * df["cagr_norm"]
        + 0.20 * df["target_market_norm"]
        + 0.20 * df["hn_norm"]
        + 0.20 * df["arxiv_norm"]
    )
    df["opportunity_score"] = (
        0.55 * df["demand_signal_score"]
        + 0.25 * (1 - df["competition_norm"])
        + 0.20 * df["fit_norm"]
    ) * 100

    df = df.sort_values("opportunity_score", ascending=False).reset_index(drop=True)
    df["priority_rank"] = df.index + 1
    return df, log_df


def build_charts(df: pd.DataFrame) -> None:
    plt.style.use("seaborn-v0_8-whitegrid")

    # Chart 1: opportunity ranking
    fig, ax = plt.subplots(figsize=(11, 6.2))
    ordered = df.sort_values("opportunity_score", ascending=True)
    ax.barh(ordered["sector"], ordered["opportunity_score"], color="#2563EB")
    ax.set_title("Opportunity Score per Sektor Jasa AI")
    ax.set_xlabel("Opportunity Score (0-100)")
    for idx, val in enumerate(ordered["opportunity_score"]):
        ax.text(val + 0.6, idx, f"{val:.1f}", va="center", fontsize=9)
    fig.tight_layout()
    fig.savefig(CHART_SCORE, dpi=220)
    plt.close(fig)

    # Chart 2: growth vs market buzz
    fig, ax = plt.subplots(figsize=(10.8, 6.2))
    bubble_size = (df["market_target_usd_billion"].pow(0.5) * 22).clip(80, 1300)
    ax.scatter(
        df["market_cagr_pct"],
        df["hn_story_hits_last_12m"],
        s=bubble_size,
        alpha=0.65,
        color="#0EA5A4",
        edgecolors="#0F172A",
        linewidths=0.6,
    )
    ax.set_title("CAGR vs Market Buzz (HN 12 Bulan)")
    ax.set_xlabel("CAGR (%)")
    ax.set_ylabel("HN Story Hits (12 bulan)")
    for _, r in df.iterrows():
        label = r["sector"].replace(" Service", "").replace(" Automation", "")
        ax.annotate(label[:28], (r["market_cagr_pct"] + 0.15, r["hn_story_hits_last_12m"] + 0.5), fontsize=8)
    fig.tight_layout()
    fig.savefig(CHART_SCATTER, dpi=220)
    plt.close(fig)

    # Chart 3: normalized signal heatmap
    cols = [
        "cagr_norm",
        "target_market_norm",
        "hn_norm",
        "arxiv_norm",
        "competition_norm",
        "fit_norm",
    ]
    heat = df.set_index("sector")[cols]
    fig, ax = plt.subplots(figsize=(11.5, 5.5))
    im = ax.imshow(heat.values, cmap="YlGnBu", aspect="auto")
    ax.set_title("Heatmap Sinyal (Ternormalisasi)")
    ax.set_xticks(range(len(cols)))
    ax.set_xticklabels(
        ["CAGR", "Target Market", "HN Buzz", "arXiv", "Competition", "Team Fit"],
        rotation=20,
        ha="right",
    )
    ax.set_yticks(range(len(heat.index)))
    ax.set_yticklabels(list(heat.index))
    for i in range(heat.shape[0]):
        for j in range(heat.shape[1]):
            ax.text(j, i, f"{heat.values[i, j]:.2f}", ha="center", va="center", fontsize=8, color="#0f172a")
    fig.colorbar(im, ax=ax, fraction=0.025, pad=0.02)
    fig.tight_layout()
    fig.savefig(CHART_HEATMAP, dpi=220)
    plt.close(fig)


def build_markdown(df: pd.DataFrame) -> None:
    top = df.iloc[0]
    generated_at = datetime.now().strftime("%d %B %Y, %H:%M")
    table_lines = [
        "| Rank | Sektor | Opportunity Score | CAGR | Market Size (Target) | GitHub Repo | HN Hits 12m |",
        "|---:|---|---:|---:|---:|---:|---:|",
    ]
    for _, r in df.iterrows():
        table_lines.append(
            f"| {int(r['priority_rank'])} | {r['sector']} | {r['opportunity_score']:.2f} | "
            f"{r['market_cagr_pct']:.2f}% | USD {r['market_target_usd_billion']:.2f}B ({int(r['market_target_year'])}) | "
            f"{int(r['github_repo_count']):,} | {int(r['hn_story_hits_last_12m']):,} |"
        )
    table_md = "\n".join(table_lines)

    source_list = "\n".join(f"- {u}" for u in df["market_source_url"].tolist())

    content = f"""# Bisnis Jasa AI Paling Potensial untuk Tim Anda

Tanggal analisis: {generated_at}
Lokasi output tetap: `/home/raymaizing/Bisnis/workspace/00-second-brain/brain-raymaizing/00-inbox`

## Ringkasan Eksekutif
- **Peluang utama saat ini:** **{top['sector']}** (Opportunity Score **{top['opportunity_score']:.2f}**).
- Tim Anda (AI enthusiast + developer + builder platform) paling cocok mendorong model **productized AI service**: kombinasi implementasi, subscription support, dan komponen platform.
- Data menunjukkan sektor dengan **CAGR tinggi** dan **market target besar** punya upside kuat, tapi perlu dijaga dengan diferensiasi agar tidak tenggelam di kompetisi repositori open-source.

## Metode Scraping (rapi dan terverifikasi)
1. Scrape market-size & CAGR dari halaman publik Precedence Research dengan `requests + BeautifulSoup`.
2. Ambil sinyal kompetisi developer dari `GitHub Search API`.
3. Ambil sinyal buzz pasar 12 bulan dari `Hacker News Algolia API`.
4. Ambil sinyal momentum riset dari `arXiv API`.
5. Semua endpoint dipanggil dengan retry (maksimal 3x) dan status/attempt dicatat di `ai_jasa_scrape_log.csv`.

## Ranking Sektor Jasa AI
{table_md}

## Rekomendasi Model Bisnis untuk Tim Anda
### 1) Fokus pada {top['sector']}
- **Kenapa sekarang:** CAGR {top['market_cagr_pct']:.2f}% dan market target USD {top['market_target_usd_billion']:.2f}B.
- **Layanan inti:** {top['recommended_service_pack']}
- **Target klien awal:** {top['target_client']}
- **Monetisasi:** {top['pricing_model']}

### 2) Paket penawaran bertingkat
- **Starter:** audit proses + quick win automation (2-4 minggu).
- **Growth:** integrasi penuh workflow + dashboard + evaluasi model.
- **Scale:** managed service bulanan dengan SLA, observability, dan continuous improvement.

### 3) GTM 90 hari (praktis)
- Minggu 1-2: pilih 1 vertical (mis. e-commerce/BPO), susun use-case template.
- Minggu 3-6: bangun 2 pilot berbayar kecil (proof value + baseline KPI).
- Minggu 7-10: productize komponen ulang-pakai (connector, prompt pipeline, evaluator).
- Minggu 11-12: scale outreach berbasis studi kasus dan metrik ROI.

## Sumber Data Utama
{source_list}
- https://api.github.com/search/repositories
- https://hn.algolia.com/api/v1/search_by_date
- https://export.arxiv.org/api/query

## Artefak yang Dihasilkan
- CSV analytics: `ai_jasa_data_analytics.csv`
- Word report + chart: `ai_jasa_data_analytics.docx`
- Chart PNG: `chart_opportunity_score.png`, `chart_cagr_vs_buzz.png`, `chart_signal_heatmap.png`
- Log scraping: `ai_jasa_scrape_log.csv`
"""
    MD_PATH.write_text(content, encoding="utf-8")


def build_docx(df: pd.DataFrame) -> None:
    generated_at = datetime.now().strftime("%d %B %Y %H:%M")
    top3 = df.head(3).copy()

    doc = Document()
    doc.add_heading("Data Analytics: Bisnis Jasa AI Paling Potensial", level=1)
    doc.add_paragraph(f"Tanggal pembuatan: {generated_at}")
    doc.add_paragraph(
        "Dokumen ini berisi hasil scraping terstruktur, ranking peluang sektor jasa AI, "
        "chart visual, dan rekomendasi implementasi untuk tim AI enthusiast + developer + builder platform."
    )

    doc.add_heading("Ringkasan Temuan", level=2)
    for _, r in top3.iterrows():
        doc.add_paragraph(
            f"Rank {int(r['priority_rank'])}: {r['sector']} (Opportunity Score {r['opportunity_score']:.2f}, "
            f"CAGR {r['market_cagr_pct']:.2f}%, market target USD {r['market_target_usd_billion']:.2f}B).",
            style="List Bullet",
        )

    doc.add_heading("Tabel Analytics Utama", level=2)
    table_cols = [
        "priority_rank",
        "sector",
        "opportunity_score",
        "market_cagr_pct",
        "market_target_usd_billion",
        "github_repo_count",
        "hn_story_hits_last_12m",
        "arxiv_total_results",
    ]
    table = doc.add_table(rows=1, cols=len(table_cols))
    table.style = "Light List Accent 1"
    header_map = {
        "priority_rank": "Rank",
        "sector": "Sektor",
        "opportunity_score": "Score",
        "market_cagr_pct": "CAGR %",
        "market_target_usd_billion": "Target Market (USD B)",
        "github_repo_count": "GitHub Repo",
        "hn_story_hits_last_12m": "HN Hits 12m",
        "arxiv_total_results": "arXiv Results",
    }
    for i, c in enumerate(table_cols):
        table.rows[0].cells[i].text = header_map[c]

    for _, r in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(int(r["priority_rank"]))
        cells[1].text = str(r["sector"])
        cells[2].text = f"{r['opportunity_score']:.2f}"
        cells[3].text = f"{r['market_cagr_pct']:.2f}"
        cells[4].text = f"{r['market_target_usd_billion']:.2f}"
        cells[5].text = f"{int(r['github_repo_count']):,}"
        cells[6].text = f"{int(r['hn_story_hits_last_12m']):,}"
        cells[7].text = f"{int(r['arxiv_total_results']):,}"

    doc.add_heading("Visualisasi Chart", level=2)
    doc.add_paragraph("1) Opportunity score per sektor")
    doc.add_picture(str(CHART_SCORE), width=Inches(6.7))
    doc.add_paragraph("2) CAGR vs market buzz (HN 12 bulan)")
    doc.add_picture(str(CHART_SCATTER), width=Inches(6.7))
    doc.add_paragraph("3) Heatmap sinyal normalisasi")
    doc.add_picture(str(CHART_HEATMAP), width=Inches(6.7))

    doc.add_heading("Aksi Implementasi 90 Hari", level=2)
    doc.add_paragraph(
        "1. Pilih 1 sektor top sebagai wedge market, siapkan paket layanan dan template delivery.",
        style="List Number",
    )
    doc.add_paragraph(
        "2. Jalankan 2 pilot berbayar kecil untuk validasi ROI, lalu ubah jadi studi kasus kuantitatif.",
        style="List Number",
    )
    doc.add_paragraph(
        "3. Productize modul reusable (connector, prompt pipeline, evaluator) untuk meningkatkan margin.",
        style="List Number",
    )
    doc.add_paragraph(
        "4. Terapkan dashboard operasional (SLA, error-rate, cost/token, business impact) untuk scale.",
        style="List Number",
    )

    doc.save(DOCX_PATH)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    df, log_df = build_dataset()

    # Persist tabular outputs first.
    numeric_round = {
        "market_base_usd_billion": 2,
        "market_target_usd_billion": 2,
        "market_cagr_pct": 2,
        "implied_cagr_pct": 2,
        "growth_multiple_x": 2,
        "demand_signal_score": 4,
        "opportunity_score": 2,
    }
    csv_df = df.copy()
    for col, digits in numeric_round.items():
        if col in csv_df.columns:
            csv_df[col] = csv_df[col].round(digits)
    csv_df.to_csv(CSV_PATH, index=False)
    log_df.to_csv(SCRAPE_LOG_PATH, index=False)

    # Build visuals and reports.
    build_charts(df)
    build_markdown(df)
    build_docx(df)

    result = {
        "csv": str(CSV_PATH),
        "markdown": str(MD_PATH),
        "docx": str(DOCX_PATH),
        "charts": [str(CHART_SCORE), str(CHART_SCATTER), str(CHART_HEATMAP)],
        "scrape_log": str(SCRAPE_LOG_PATH),
        "rows": len(df),
    }
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()
